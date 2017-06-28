#!/usr/bin/env python
# -*- coding: utf-8 -*-
# Created by xiaoqin00 on 2017/6/26

from optparse import OptionParser
from docx import Document
from docx.shared import Inches
def convert():
    #创建 Document 对象，相当于打开一个 word 文档
    document = Document()

    # #向文档中添加一个标题，标题级别设置为0级
    # document.add_heading('This is title', level=0)

    #向文档中添加一个段落，并将段落引用赋给变量 p
    #使用 add_run 方法追加字段，并设置格式
    # f=open('test.pdf.txt','r')
    f=open(options.input,'r')
    for i in f.readlines():
        print i
        print 'test'
        i=str(i)
        i=i.split()
        if not i:
            i='\t'
            # continue
        p = document.add_paragraph(i)

    # p.add_run('bold').bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True
    #
    # #添加标题和段落，采用不同的形式
    # document.add_heading('This is Heading, level 1', level=1)
    # document.add_paragraph('Intese quote',style="Intense Quote")
    # document.add_paragraph('first item in unordered list', style='List Bullet')
    # document.add_paragraph('first item in ordered list', style='List Number')
    #
    # # #添加图片，设置图片大小
    # # document.add_picture(r"D:\picture\a.jpg", width=Inches(2.25))
    #
    # #添加表格，填入表格内容
    # table = document.add_table(rows=2, cols=2)
    # table.cell(0,0).text = "cell_00"
    # table.cell(0,1).text = "cell_01"
    # table.cell(1,0).text = "cell_10"
    # table.cell(1,1).text = "cell_11"

    #保存文本
    if options.output:
        document.save(options.output)
    document.save('demo.docx')
    return

if __name__ == '__main__':
    parser=OptionParser(usage='%prog [options]')
    parser.add_option('-i','--in',dest='input',help='input file')
    parser.add_option('-o','--out',dest='output',help='output file')
    (options,args)=parser.parse_args()
    # print options.input
    convert()