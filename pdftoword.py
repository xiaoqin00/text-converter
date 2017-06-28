#!/usr/bin/env python
# -*- coding: utf-8 -*-
# Created by xiaoqin00 on 2017/6/26

#pdf 转为word,没有找到pdf直接转换为word的方法，就先转为txt，然后转换为word

import sys
from pdfminer.pdfinterp import PDFResourceManager,PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage

from optparse import OptionParser
from docx import Document
from docx.shared import Inches

#main
def pdftotxt():
    #输出文件名，这里只处理单文档，所以只用了argv［1］
    outfile = options.input + '.txt'
    args = [options.input]

    debug = 0
    pagenos = set()
    password = ''
    maxpages = 0
    rotation = 0
    codec = 'utf-8'   #输出编码
    caching = True
    imagewriter = None
    laparams = LAParams()
    #
    PDFResourceManager.debug = debug
    PDFPageInterpreter.debug = debug

    rsrcmgr = PDFResourceManager(caching=caching)
    outfp = file(outfile,'w')
#pdf转换
    device = TextConverter(rsrcmgr, outfp, codec=codec, laparams=laparams,
                imagewriter=imagewriter)

    for fname in args:
        fp = file(fname,'rb')
        interpreter = PDFPageInterpreter(rsrcmgr, device)
#处理文档对象中每一页的内容
        for page in PDFPage.get_pages(fp, pagenos,
                          maxpages=maxpages, password=password,
                          caching=caching, check_extractable=True) :
            page.rotate = (page.rotate+rotation) % 360
            interpreter.process_page(page)
        fp.close()
    device.close()
    outfp.close()
    return

def txttoword():
    #创建 Document 对象，相当于打开一个 word 文档
    document = Document()

    # #向文档中添加一个标题，标题级别设置为0级
    # document.add_heading('This is title', level=0)

    #向文档中添加一个段落，并将段落引用赋给变量 p
    #使用 add_run 方法追加字段，并设置格式
    # f=open('test.pdf.txt','r')
    f=open(options.input+'.txt','r')
    for i in f.readlines():
        print i
        print 'test'
        i=str(i)
        i=i.split()
        if not i:
            i='\t'
            # continue
        p = document.add_paragraph(i)

    # p.add_run('bold').bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True
    #
    # #添加标题和段落，采用不同的形式
    # document.add_heading('This is Heading, level 1', level=1)
    # document.add_paragraph('Intese quote',style="Intense Quote")
    # document.add_paragraph('first item in unordered list', style='List Bullet')
    # document.add_paragraph('first item in ordered list', style='List Number')
    #
    # # #添加图片，设置图片大小
    # # document.add_picture(r"D:\picture\a.jpg", width=Inches(2.25))
    #
    # #添加表格，填入表格内容
    # table = document.add_table(rows=2, cols=2)
    # table.cell(0,0).text = "cell_00"
    # table.cell(0,1).text = "cell_01"
    # table.cell(1,0).text = "cell_10"
    # table.cell(1,1).text = "cell_11"

    #保存文本
    if options.output:
        document.save(options.output)
    document.save(options.input+'.docx')
    return

if __name__ == '__main__':
    parser=OptionParser(usage='%prog [options]')
    parser.add_option('-i','--in',dest='input',help='input file')
    parser.add_option('-o','--out',dest='output',help='output file')
    (options,args)=parser.parse_args()
    # print options.input
    pdftotxt()
    txttoword()